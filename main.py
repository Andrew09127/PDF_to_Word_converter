import PyPDF2
import pymupdf as fitz
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import io
from PIL import Image
import os
import shutil
from pathlib import Path
import time
import logging
from datetime import datetime
import gc
import threading
from queue import Queue
import json

os.environ.setdefault("FLAGS_use_mkldnn", "0")
os.environ.setdefault("FLAGS_use_onednn", "0")

try:
    from pdf2docx import Converter
except ImportError:
    Converter = None

try:
    from docxcompose.composer import Composer
except ImportError:
    Composer = None

try:
    import cv2
    import numpy as np
    from paddleocr import PaddleOCR, PPStructure
    from paddleocr.ppstructure.recovery.recovery_to_doc import (
        convert_info_docx,
        sorted_layout_boxes,
    )
except ImportError:
    cv2 = None
    np = None
    PaddleOCR = None
    PPStructure = None
    convert_info_docx = None
    sorted_layout_boxes = None

PADDLE_LANG = os.getenv("PADDLE_LANG", "en")
PADDLE_OCR_LANG = os.getenv("PADDLE_OCR_LANG", "ru")
OCR_DPI = int(os.getenv("OCR_DPI", "300"))
OCR_GPU = os.getenv("OCR_GPU", "0").lower() in ("1", "true", "yes")
PADDLE_CPU_THREADS = int(os.getenv("PADDLE_CPU_THREADS", "4"))
PADDLE_FALLBACK_TO_OCR = os.getenv("PADDLE_FALLBACK_TO_OCR", "1").lower() in ("1", "true", "yes")
CONVERT_MODE = os.getenv("CONVERT_MODE", "auto").lower()

# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('pdf_conversion.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)

class PDFPipelineConverter:
    def __init__(self, input_folder, output_folder, backup_folder=None):
        """
        input_folder: папка с PDF файлами (будут удаляться после конвертации)
        output_folder: папка для сохранения DOCX файлов
        backup_folder: опциональная папка для бэкапа PDF перед удалением
        """
        self.input_folder = Path(input_folder)
        self.output_folder = Path(output_folder)
        self.backup_folder = Path(backup_folder) if backup_folder else None
        self.paddle_structure_engine = None
        self.paddle_ocr_engine = None
        
        # Статистика
        self.stats = {
            'processed': 0,
            'success': 0,
            'failed': 0,
            'skipped': 0,
            'total_size_saved_mb': 0,
            'start_time': None,
            'end_time': None
        }
        
        # Создаем папки
        self.output_folder.mkdir(parents=True, exist_ok=True)
        if self.backup_folder:
            self.backup_folder.mkdir(parents=True, exist_ok=True)
        
        # Файл для сохранения состояния
        self.checkpoint_file = self.input_folder / 'conversion_state.json'
        
    def extract_text_from_pdf(self, pdf_path):
        """Извлечение текста из PDF"""
        text_content = []
        try:
            with open(pdf_path, "rb") as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text:
                            text_content.append(f"--- Page {page_num + 1} ---\n{text}\n")
                    except Exception as e:
                        text_content.append(f"--- Page {page_num + 1} ---\n[Error: {str(e)}]\n")
            return text_content
        except Exception as e:
            logging.error(f"Text extraction failed: {e}")
            return [f"Error extracting text: {str(e)}"]
    
    def extract_images_from_pdf(self, pdf_path):
        """Извлечение изображений из PDF"""
        images_content = []
        pdf_document = None
        try:
            pdf_document = fitz.open(pdf_path)
            for page_num in range(min(pdf_document.page_count, 100)):  # Ограничение на 100 страниц
                page = pdf_document[page_num]
                page_images = []
                
                # Получаем изображения со страницы
                image_list = page.get_images(full=True)
                for img_index, img in enumerate(image_list[:5]):  # Не более 5 изображений на страницу
                    try:
                        xref = img[0]
                        base_image = pdf_document.extract_image(xref)
                        img_bytes = base_image["image"]
                        
                        # Конвертируем в PNG
                        with io.BytesIO(img_bytes) as img_stream:
                            with Image.open(img_stream) as pil_image:
                                # Уменьшаем размер изображения если оно слишком большое
                                if pil_image.size[0] > 1000 or pil_image.size[1] > 1000:
                                    pil_image.thumbnail((800, 800), Image.Resampling.LANCZOS)
                                
                                png_stream = io.BytesIO()
                                pil_image.save(png_stream, format="PNG", optimize=True)
                                png_stream.seek(0)
                                page_images.append(png_stream)
                                
                    except Exception as e:
                        logging.warning(f"Failed to extract image on page {page_num}: {e}")
                        continue
                
                images_content.append(page_images)
                if page_num % 20 == 0:
                    gc.collect()  # Периодическая очистка
                    
        except Exception as e:
            logging.error(f"Image extraction failed: {e}")
        finally:
            if pdf_document:
                pdf_document.close()
        
        return images_content

    def is_scanned_pdf(self, pdf_path, min_text_chars=40, pages_to_check=3):
        """Проверяет, есть ли в PDF текстовый слой."""
        pdf_document = None
        try:
            pdf_document = fitz.open(pdf_path)
            text = []
            for page_index in range(min(pdf_document.page_count, pages_to_check)):
                page = pdf_document[page_index]
                text.append(page.get_text("text") or "")
            return len("".join(text).strip()) < min_text_chars
        except Exception as e:
            logging.warning(f"Failed to detect PDF type for {pdf_path.name}: {e}")
            return False
        finally:
            if pdf_document:
                pdf_document.close()

    def convert_scanned_pdf_to_docx(self, pdf_path, docx_path):
        """OCR-конвертация сканированного PDF в DOCX через PaddleOCR PP-Structure."""
        if (
            cv2 is None
            or np is None
            or PPStructure is None
            or sorted_layout_boxes is None
            or convert_info_docx is None
        ):
            raise ImportError(
                "Не установлен PaddleOCR/PP-Structure. Установите зависимости из README."
            )

        try:
            self.convert_scanned_pdf_with_structure(pdf_path, docx_path)
        except Exception as e:
            if not PADDLE_FALLBACK_TO_OCR:
                raise
            logging.warning(
                f"PP-Structure failed for {pdf_path.name}; falling back to PaddleOCR text OCR: {e}"
            )
            self.convert_scanned_pdf_with_paddle_ocr(pdf_path, docx_path)

    def convert_scanned_pdf_with_structure(self, pdf_path, docx_path):
        """Пытается восстановить структуру документа через PaddleOCR PP-Structure."""
        pdf_document = None
        try:
            if self.paddle_structure_engine is None:
                logging.info(
                    f"Initializing PaddleOCR PP-Structure: lang={PADDLE_LANG}, gpu={OCR_GPU}"
                )
                self.paddle_structure_engine = PPStructure(
                    recovery=True,
                    lang=PADDLE_LANG,
                    use_gpu=OCR_GPU,
                    ir_optim=False,
                    enable_mkldnn=False,
                    cpu_threads=PADDLE_CPU_THREADS,
                    show_log=True,
                )

            pdf_document = fitz.open(pdf_path)
            page_docx_paths = []
            temp_folder = self.output_folder / "_paddle_tmp" / pdf_path.stem
            temp_folder.mkdir(parents=True, exist_ok=True)
            zoom = OCR_DPI / 72
            matrix = fitz.Matrix(zoom, zoom)

            for page_num, page in enumerate(pdf_document, 1):
                logging.info(
                    f"PaddleOCR page {page_num}/{pdf_document.page_count}: {pdf_path.name}"
                )
                pixmap = page.get_pixmap(matrix=matrix, alpha=False)
                image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
                image_bgr = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)

                result = self.paddle_structure_engine(image_bgr)
                result_for_docx = sorted_layout_boxes(result, image_bgr.shape[1])
                page_name = f"page_{page_num:04d}"
                convert_info_docx(
                    image_bgr,
                    result_for_docx,
                    str(temp_folder),
                    page_name,
                )
                page_docx_path = temp_folder / f"{page_name}.docx"
                page_docx_paths.append(page_docx_path)

                if page_num % 10 == 0:
                    gc.collect()

            self.merge_docx_pages(page_docx_paths, docx_path)
        finally:
            if pdf_document:
                pdf_document.close()

    def convert_scanned_pdf_with_paddle_ocr(self, pdf_path, docx_path):
        """Fallback: распознает текст PaddleOCR и собирает редактируемый DOCX по координатам."""
        if cv2 is None or np is None or PaddleOCR is None:
            raise ImportError("Не установлен PaddleOCR. Установите зависимости из README.")

        pdf_document = None
        try:
            if self.paddle_ocr_engine is None:
                logging.info(
                    f"Initializing PaddleOCR text OCR: lang={PADDLE_OCR_LANG}, gpu={OCR_GPU}"
                )
                self.paddle_ocr_engine = PaddleOCR(
                    lang=PADDLE_OCR_LANG,
                    use_gpu=OCR_GPU,
                    ir_optim=False,
                    enable_mkldnn=False,
                    cpu_threads=PADDLE_CPU_THREADS,
                    show_log=True,
                )

            pdf_document = fitz.open(pdf_path)
            doc = Document()
            zoom = OCR_DPI / 72
            matrix = fitz.Matrix(zoom, zoom)

            for page_num, page in enumerate(pdf_document, 1):
                logging.info(
                    f"PaddleOCR text page {page_num}/{pdf_document.page_count}: {pdf_path.name}"
                )
                pixmap = page.get_pixmap(matrix=matrix, alpha=False)
                image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
                image_bgr = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)

                result = self.paddle_ocr_engine.ocr(image_bgr, cls=False)
                layout_lines = self.build_layout_lines_from_paddle_result(
                    result,
                    image_bgr.shape[1],
                )

                if page_num > 1:
                    doc.add_page_break()
                self.add_ocr_layout_to_doc(doc, layout_lines, image_bgr.shape[1])

            doc.save(docx_path)
        finally:
            if pdf_document:
                pdf_document.close()

    def build_layout_lines_from_paddle_result(self, result, page_width_px):
        """Собирает результат PaddleOCR в строки с координатами."""
        blocks = []
        page_results = result[0] if result and isinstance(result[0], list) else result
        for item in page_results or []:
            if len(item) < 2:
                continue
            box, text_info = item[0], item[1]
            text = self.clean_docx_text(str(text_info[0]).strip())
            confidence = float(text_info[1]) if len(text_info) > 1 else 1.0
            if not text or confidence < 0.35:
                continue

            xs = [point[0] for point in box]
            ys = [point[1] for point in box]
            x1, x2 = min(xs), max(xs)
            y1, y2 = min(ys), max(ys)
            height = max(y2 - y1, 1)
            blocks.append({
                "text": text,
                "x1": x1,
                "x2": x2,
                "y1": y1,
                "y2": y2,
                "height": height,
            })

        if not blocks:
            return []

        blocks.sort(key=lambda block: (block["y1"], block["x1"]))
        lines = []
        for block in blocks:
            center_y = (block["y1"] + block["y2"]) / 2
            matched_line = None
            for line in lines:
                tolerance = max(line["height"], block["height"]) * 0.65
                if abs(center_y - line["center_y"]) <= tolerance:
                    matched_line = line
                    break

            if matched_line is None:
                lines.append({
                    "center_y": center_y,
                    "height": block["height"],
                    "blocks": [block],
                })
            else:
                matched_line["blocks"].append(block)
                matched_line["height"] = max(matched_line["height"], block["height"])
                matched_line["center_y"] = (
                    matched_line["center_y"] * (len(matched_line["blocks"]) - 1) + center_y
                ) / len(matched_line["blocks"])

        lines.sort(key=lambda line: line["center_y"])
        for line in lines:
            line["blocks"].sort(key=lambda block: block["x1"])
            line["x1"] = min(block["x1"] for block in line["blocks"])
            line["y1"] = min(block["y1"] for block in line["blocks"])
            line["y2"] = max(block["y2"] for block in line["blocks"])
            line["height"] = max(block["height"] for block in line["blocks"])
            line["page_width_px"] = page_width_px
        return lines

    def clean_docx_text(self, text):
        """Удаляет символы, которые Word/docx не принимает в XML."""
        return "".join(
            char for char in text
            if char in ("\n", "\t") or ord(char) >= 32
        )

    def add_ocr_layout_to_doc(self, doc, layout_lines, page_width_px):
        """Добавляет OCR-текст в DOCX с приближенными размерами и отступами."""
        if not layout_lines:
            doc.add_paragraph("[OCR text not detected on this page]")
            return

        section = doc.sections[-1]
        usable_width_twips = section.page_width - section.left_margin - section.right_margin
        previous_y2 = None

        for line in layout_lines:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.left_indent = int(
                usable_width_twips * (line["x1"] / max(page_width_px, 1))
            )

            if previous_y2 is not None:
                gap_px = max(line["y1"] - previous_y2, 0)
                paragraph.paragraph_format.space_before = Pt(min(gap_px * 72 / OCR_DPI, 36))
            else:
                paragraph.paragraph_format.space_before = Pt(0)

            previous_x2 = line["x1"]
            for index, block in enumerate(line["blocks"]):
                if index > 0:
                    gap_ratio = max((block["x1"] - previous_x2) / max(page_width_px, 1), 0)
                    spaces = max(1, min(int(gap_ratio * 130), 20))
                    paragraph.add_run(" " * spaces)

                run = paragraph.add_run(block["text"])
                run.font.name = "Arial"
                font_size = block["height"] * 72 / OCR_DPI * 0.85
                run.font.size = Pt(max(7, min(font_size, 22)))
                previous_x2 = block["x2"]

            previous_y2 = line["y2"]

    def merge_docx_pages(self, page_docx_paths, output_docx_path):
        """Объединяет DOCX-страницы PaddleOCR в один DOCX."""
        if not page_docx_paths:
            raise ValueError("PaddleOCR did not create any DOCX pages")

        if len(page_docx_paths) == 1:
            shutil.copyfile(page_docx_paths[0], output_docx_path)
            return

        if Composer is not None:
            merged_doc = Document(page_docx_paths[0])
            composer = Composer(merged_doc)
            for page_docx_path in page_docx_paths[1:]:
                composer.append(Document(page_docx_path))
            composer.save(output_docx_path)
            return

        merged_doc = Document(page_docx_paths[0])
        for page_docx_path in page_docx_paths[1:]:
            merged_doc.add_page_break()
            page_doc = Document(page_docx_path)
            for element in page_doc.element.body:
                if element.tag.endswith("sectPr"):
                    continue
                merged_doc.element.body.append(element)

        merged_doc.save(output_docx_path)
    
    def convert_single_pdf(self, pdf_path, docx_path):
        """Конвертация одного PDF в редактируемый DOCX"""
        converter = None
        try:
            use_paddleocr = CONVERT_MODE == "paddleocr" or (
                CONVERT_MODE == "auto" and self.is_scanned_pdf(pdf_path)
            )

            if use_paddleocr:
                logging.info(f"Scanned PDF detected, using PaddleOCR: {pdf_path.name}")
                self.convert_scanned_pdf_to_docx(pdf_path, docx_path)
            else:
                if CONVERT_MODE not in ("auto", "pdf2docx"):
                    raise ValueError("CONVERT_MODE must be one of: auto, paddleocr, pdf2docx")

                if Converter is None:
                    raise ImportError(
                        "Не установлена библиотека pdf2docx. Установите ее командой: python -m pip install pdf2docx"
                    )

                converter = Converter(str(pdf_path))
                converter.convert(str(docx_path), start=0, end=None)

            # Получаем размер файла для статистики
            file_size = os.path.getsize(pdf_path) / (1024 * 1024)
            
            return True, file_size
            
        except MemoryError:
            logging.error(f"Memory error while converting {pdf_path.name}")
            return False, 0
        except Exception as e:
            logging.error(f"Conversion error for {pdf_path.name}: {str(e)}")
            return False, 0
        finally:
            if converter:
                converter.close()
            gc.collect()
    
    def safe_delete_pdf(self, pdf_path, move_to_backup=True):
        """Безопасное удаление PDF файла (с возможностью перемещения в бэкап)"""
        try:
            if move_to_backup and self.backup_folder:
                # Перемещаем в папку бэкапа вместо удаления
                backup_path = self.backup_folder / pdf_path.name
                shutil.move(str(pdf_path), str(backup_path))
                logging.debug(f"Moved to backup: {pdf_path.name}")
            else:
                # Удаляем файл
                os.remove(pdf_path)
                logging.debug(f"Deleted: {pdf_path.name}")
            return True
        except Exception as e:
            logging.error(f"Failed to delete/move {pdf_path.name}: {e}")
            return False
    
    def save_checkpoint(self):
        """Сохранение состояния обработки"""
        checkpoint_data = {
            'processed': self.stats['processed'],
            'success': self.stats['success'],
            'failed': self.stats['failed'],
            'total_size_saved_mb': self.stats['total_size_saved_mb'],
            'last_update': datetime.now().isoformat()
        }
        try:
            with open(self.checkpoint_file, 'w') as f:
                json.dump(checkpoint_data, f, indent=2)
        except Exception as e:
            logging.warning(f"Failed to save checkpoint: {e}")
    
    def load_checkpoint(self):
        """Загрузка состояния обработки"""
        if self.checkpoint_file.exists():
            try:
                with open(self.checkpoint_file, 'r') as f:
                    data = json.load(f)
                    self.stats.update(data)
                logging.info(f"Loaded checkpoint: {self.stats['success']} files already processed")
                return True
            except Exception as e:
                logging.warning(f"Failed to load checkpoint: {e}")
        return False
    
    def process_pipeline(self, delete_after_processing=True, move_to_backup=False):
        """
        Основной конвейер обработки
        delete_after_processing: удалять PDF после конвертации
        move_to_backup: перемещать PDF в папку бэкапа перед удалением
        """
        self.stats['start_time'] = time.time()
        
        # Загружаем чекпоинт
        self.load_checkpoint()
        
        # Получаем список PDF файлов
        pdf_files = sorted(self.input_folder.glob("*.pdf"))
        
        # Фильтруем уже обработанные (если есть чекпоинт)
        total_found = len(pdf_files)
        logging.info(f"Found {total_found} PDF files in {self.input_folder}")
        
        # Мониторинг дискового пространства
        def check_disk_space(folder, required_gb=1):
            free_space = shutil.disk_usage(folder).free / (1024**3)
            if free_space < required_gb:
                logging.warning(f"Low disk space on {folder}: only {free_space:.1f} GB free")
                return False
            return True
        
        # Основной цикл обработки
        for i, pdf_path in enumerate(pdf_files, 1):
            # Проверяем, не обработан ли уже файл (по имени)
            docx_path = self.output_folder / f"{pdf_path.stem}.docx"
            if docx_path.exists():
                logging.info(f"[{i}/{total_found}] Skipping (already converted): {pdf_path.name}")
                self.stats['skipped'] += 1
                if delete_after_processing:
                    self.safe_delete_pdf(pdf_path, move_to_backup=move_to_backup)
                continue
            
            logging.info(f"[{i}/{total_found}] Processing: {pdf_path.name} ({pdf_path.stat().st_size / (1024*1024):.2f} MB)")
            
            # Конвертируем PDF в DOCX
            success, file_size = self.convert_single_pdf(pdf_path, docx_path)
            
            if success:
                self.stats['success'] += 1
                self.stats['total_size_saved_mb'] += file_size
                
                # Удаляем или перемещаем исходный PDF
                if delete_after_processing:
                    if self.safe_delete_pdf(pdf_path, move_to_backup=move_to_backup):
                        logging.info(f"✓ Converted and removed: {pdf_path.name}")
                    else:
                        logging.warning(f"✓ Converted but failed to remove: {pdf_path.name}")
                else:
                    logging.info(f"✓ Converted successfully: {pdf_path.name}")
            else:
                self.stats['failed'] += 1
                logging.error(f"✗ Failed to convert: {pdf_path.name}")
            
            self.stats['processed'] = i
            
            # Каждые 10 файлов сохраняем чекпоинт и выводим статистику
            if i % 10 == 0:
                self.save_checkpoint()
                self.print_stats(i, total_found)
                check_disk_space(self.output_folder, 0.5)
            
            # Очищаем память каждые 50 файлов
            if i % 50 == 0:
                gc.collect()
                logging.info("Memory cleanup performed")
        
        self.stats['end_time'] = time.time()
        self.print_final_stats()
        
        # Удаляем чекпоинт после успешного завершения
        if self.stats['failed'] == 0:
            try:
                self.checkpoint_file.unlink()
                logging.info("Checkpoint file removed (all files processed successfully)")
            except:
                pass
    
    def print_stats(self, current, total):
        """Вывод промежуточной статистики"""
        elapsed = time.time() - self.stats['start_time']
        rate = current / elapsed if elapsed > 0 else 0
        remaining_files = total - current
        eta = remaining_files / rate if rate > 0 else 0
        
        logging.info(f"📊 Progress: {current}/{total} | "
                    f"✅ {self.stats['success']} | "
                    f"❌ {self.stats['failed']} | "
                    f"⏭️ {self.stats['skipped']} | "
                    f"📁 Saved: {self.stats['total_size_saved_mb']:.1f} MB | "
                    f"⚡ {rate:.2f} files/sec | "
                    f"⏱️ ETA: {eta/60:.1f} min")
    
    def print_final_stats(self):
        """Вывод финальной статистики"""
        total_time = self.stats['end_time'] - self.stats['start_time']
        logging.info("\n" + "="*60)
        logging.info("🎉 CONVERSION PIPELINE COMPLETED 🎉")
        logging.info("="*60)
        logging.info(f"⏱️  Total time: {total_time/60:.1f} minutes ({total_time/3600:.2f} hours)")
        logging.info(f"✅ Successfully converted: {self.stats['success']}")
        logging.info(f"❌ Failed: {self.stats['failed']}")
        logging.info(f"⏭️  Skipped: {self.stats['skipped']}")
        logging.info(f"💾 Total disk space saved: {self.stats['total_size_saved_mb']:.1f} MB")
        logging.info(f"📊 Average speed: {self.stats['success']/(total_time/3600):.1f} files/hour")
        
        if self.stats['failed'] > 0:
            logging.warning(f"⚠️  {self.stats['failed']} files failed. Check log for details.")
        else:
            logging.info("✨ All files processed successfully! ✨")
        logging.info("="*60)

# Запуск конвейера
def run_conversion_pipeline():
    """
    Пример настройки и запуска конвейера
    """
    # Настройка путей
    INPUT_FOLDER = "./pdf_to_convert"      # Папка с PDF (файлы будут удаляться)
    OUTPUT_FOLDER = "./converted_word"     # Папка для DOCX файлов
    
    # Создаем экземпляр конвертера
    converter = PDFPipelineConverter(
        input_folder=INPUT_FOLDER,
        output_folder=OUTPUT_FOLDER,
        backup_folder=None
    )
    
    # Запускаем обработку
    # delete_after_processing=True - удалять PDF после конвертации
    # move_to_backup=False - удалять PDF сразу, без папки бэкапа
    converter.process_pipeline(
        delete_after_processing=True,  # Удалять оригиналы
        move_to_backup=False            # Удалять сразу, без бэкапа
    )

# Альтернативный простой вариант (минимальная настройка)
def simple_pipeline():
    """Простой вариант - только конвертация и удаление без бэкапов"""
    
    INPUT_FOLDER = "./pdf_files"
    OUTPUT_FOLDER = "./word_files"
    
    converter = PDFPipelineConverter(
        input_folder=INPUT_FOLDER,
        output_folder=OUTPUT_FOLDER,
        backup_folder=None  # Без бэкапа - сразу удаляем
    )
    
    converter.process_pipeline(
        delete_after_processing=True,
        move_to_backup=False  # Просто удаляем
    )

if __name__ == "__main__":
    # Выберите нужный вариант
    
    # Вариант 1: Полный с бэкапом (рекомендуется для важных файлов)
    run_conversion_pipeline()
    
    # Вариант 2: Простой без бэкапа (максимальная экономия места)
    # simple_pipeline()
